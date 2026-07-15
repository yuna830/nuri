from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("woori_link_recall_action_guide.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True, color="1F3A5F")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.color.rgb = RGBColor(46, 116, 181)
    run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Malgun Gothic"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(10)
    doc.add_paragraph()


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def main():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Woori Link 리콜 조치 요청 연동 가이드")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run("어르신 앱 → DB → 보호자/복지사 화면 전달 구조")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(90, 100, 115)

    add_callout(
        doc,
        "핵심 요약",
        "어르신 앱에서 리콜 조치 요청을 누르면 wl_action_records 테이블에 action_type=RECALL, "
        "action_status=PENDING 상태의 조치 기록이 생성됩니다. 보호자/복지사 화면은 이 기록을 조회해 "
        "'리콜 미조치 대상자'로 표시하면 됩니다.",
    )

    add_heading(doc, "1. 저장 위치", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["DB 테이블", "wl_action_records"],
            ["기록 의미", "어르신이 요청한 리콜 조치/방문/전화 안내 기록"],
            ["리콜 요청 구분", "action_type = RECALL"],
            ["초기 상태", "action_status = PENDING"],
        ],
        [Inches(1.7), Inches(4.8)],
    )

    add_heading(doc, "2. 저장되는 주요 컬럼", 1)
    add_table(
        doc,
        ["컬럼", "예시 값", "설명"],
        [
            ["id", "6", "조치 요청 고유 ID"],
            ["senior_id", "3", "요청한 어르신 ID"],
            ["welfare_worker_id", "null", "담당 복지사 ID. 현재는 null 가능"],
            ["action_type", "RECALL", "리콜 조치 요청 타입"],
            ["action_subject", "SENIOR", "요청 주체"],
            ["action_status", "PENDING", "미조치/진행/완료 상태"],
            ["product_name", "전기요", "리콜 대상 제품명"],
            ["note", "요청 메모 + 모델명 + 리콜 사유", "보호자/복지사가 확인할 상세 내용"],
            ["created_at", "2026-07-14T16:20:50", "요청 생성 시간"],
            ["updated_at", "2026-07-14T16:20:50", "상태 변경 시간"],
        ],
        [Inches(1.4), Inches(1.7), Inches(3.4)],
    )

    add_heading(doc, "3. Flutter에서 생성하는 요청 Payload", 1)
    add_table(
        doc,
        ["필드", "값"],
        [
            ["seniorId", "현재 로그인한 어르신 ID"],
            ["actionType", "RECALL"],
            ["actionSubject", "SENIOR"],
            ["status", "PENDING"],
            ["productName", "리콜 제품명"],
            ["note", "제가 가진 제품이 리콜 대상이라고 확인되었습니다.\\n모델명: DRC-LS0421\\n보호자나 복지사에게 전화 또는 방문 안내를 요청합니다.\\n\\n제품안전정보센터 리콜 사유: ..."],
        ],
        [Inches(1.6), Inches(4.9)],
    )

    add_heading(doc, "4. 보호자/복지사 화면에서 받아야 할 값", 1)
    add_bullet(doc, "id: 조치 요청 ID. 상태 변경 API 호출 시 사용합니다.")
    add_bullet(doc, "seniorId: 어떤 어르신의 요청인지 식별합니다.")
    add_bullet(doc, "actionType: RECALL인 경우 리콜 조치 요청으로 표시합니다.")
    add_bullet(doc, "status: PENDING이면 '리콜 미조치 대상자'로 표시합니다.")
    add_bullet(doc, "productName: 리콜 대상 제품명입니다.")
    add_bullet(doc, "note: 모델명, 요청 메모, 제품안전정보센터 리콜 사유가 들어갑니다.")
    add_bullet(doc, "createdAt / updatedAt: 요청일과 상태 변경일 표시용입니다.")

    add_heading(doc, "5. 조회 및 상태 변경 API", 1)
    add_table(
        doc,
        ["용도", "API"],
        [
            ["전체 미조치/대기 조치 목록", "GET /api/actions/pending"],
            ["어르신별 조치 목록", "GET /api/actions/senior/{seniorId}"],
            ["복지사별 조치 목록", "GET /api/actions/welfare-worker/{welfareWorkerId}"],
            ["조치 상태 변경", "PATCH /api/actions/{id}/status?status=COMPLETED"],
        ],
        [Inches(2.2), Inches(4.3)],
    )

    add_heading(doc, "6. 상태값 매핑", 1)
    add_table(
        doc,
        ["DB/API 상태", "화면 표시", "의미"],
        [
            ["PENDING", "미조치", "요청은 생성됐지만 아직 연락/방문 조치 전"],
            ["IN_PROGRESS", "조치 중", "보호자/복지사가 연락 또는 방문 조치를 진행 중"],
            ["COMPLETED", "조치 완료", "안내/방문/교환/환불 등 조치 완료"],
            ["CANCELLED", "취소", "요청 취소 또는 조치 불필요 처리"],
        ],
        [Inches(1.5), Inches(1.4), Inches(3.6)],
    )

    add_heading(doc, "7. 보호자/복지사 화면 구현 기준", 1)
    add_callout(
        doc,
        "필터 조건",
        "actionType == RECALL 이고 status == PENDING 인 항목을 '리콜 미조치 대상자'로 표시합니다.",
    )
    add_bullet(doc, "목록 카드 예시: 제품명, 어르신명/ID, 요청일, 상태(미조치)를 표시합니다.")
    add_bullet(doc, "상세 화면 예시: 모델명, 요청 메모, 리콜 사유, 문의처를 표시합니다.")
    add_bullet(doc, "조치 후에는 PATCH API로 status를 IN_PROGRESS 또는 COMPLETED로 변경합니다.")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = footer.add_run("작성 기준: Woori Link Flutter/Spring 리콜 조치 요청 연동")
    rr.font.name = "Malgun Gothic"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor(120, 128, 140)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
